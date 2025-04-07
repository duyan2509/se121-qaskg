import os
import json
from typing import List
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter



class DocxTextExtractor:
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap


    def escape_latex_chars(self, text: str) -> str:
        """Escape special LaTeX characters."""
        # Các ký tự cần escape trong LaTeX
        escape_chars = {
            '&': '\\&',
            '%': '\\%',
            '$': '\\$',
            '#': '\\#',
            '_': '\\_',
            '{': '\\{',
            '}': '\\}',
            '~': '\\textasciitilde{}',
            '^': '\\textasciicircum{}',
            '\\': '\\textbackslash{}'
        }

        for char, escape in escape_chars.items():
            text = text.replace(char, escape)

        return text


    def table_to_latex(self, table) -> str:
        """Convert a docx table to LaTeX format."""
        rows = len(table.rows)
        cols = len(table.columns)

        # Start LaTeX table
        latex = "\\begin{table}[h]\n\\centering\n"

        # Create column format (all centered)
        latex += "\\begin{tabular}{" + "|c" * cols + "|}\n\\hline\n"

        # Process table rows
        for i, row in enumerate(table.rows):
            # Get text from each cell in the row and process it
            cells = [self.escape_latex_chars(cell.text.strip()) for cell in row.cells]

            # Add row to LaTeX table
            latex += " & ".join(cells) + " \\\\\n\\hline\n"

        # Close LaTeX table
        latex += "\\end{tabular}\n\\end{table}\n"

        return latex

    def extract_content(self, docx_path: str) -> List[str]:
        """Extract text and tables from docx, marking each segment if it's a table."""
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"File '{docx_path}' not found!")

        doc = Document(docx_path)
        for paragraph in doc.paragraphs:
            print(repr(paragraph.text))
            # Extract header from first section
        header = ""
        if doc.sections and doc.sections[0].header:
            header_paras = doc.sections[0].header.paragraphs
            header = "\n".join(p.text.strip() for p in header_paras if p.text.strip())

        # This will hold text content and tables in order of appearance
        content_elements = []

        # First, collect header if any
        if header:
            content_elements.append((header, False))

        # Process all paragraphs and tables
        for element in doc.element.body:
            if element.tag.endswith('tbl'):
                # It's a table - convert to LaTeX
                table = doc.tables[len([e for e in content_elements if e[1]])]  # Count previous tables
                latex_table = self.table_to_latex(table)
                content_elements.append((latex_table, True))
            elif element.tag.endswith('p'):
                # It's a paragraph
                for paragraph in doc.paragraphs:
                    if paragraph._p == element:  # Match the paragraph object
                        if paragraph.text.strip():
                            content_elements.append((paragraph.text.strip(), False))
                        break

        # Join all text elements, keeping track of table positions
        full_text = ""
        table_positions = []
        current_position = 0

        for content, is_table in content_elements:
            if is_table:
                table_positions.append((current_position, content))
            else:
                full_text += content + "\n\n"
                current_position = len(full_text)

        # Chunking the text content
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap
        )
        chunks = text_splitter.split_text(full_text)

        # Re-insert tables into the appropriate chunks
        final_chunks = []
        for chunk in chunks:
            chunk_start = full_text.find(chunk)
            if chunk_start == -1:  # Handle potential edge cases with overlapping chunks
                final_chunks.append(chunk)
                continue

            chunk_end = chunk_start + len(chunk)

            # Find tables that belong in this chunk
            tables_in_chunk = []
            for table_pos, table_content in table_positions:
                if chunk_start <= table_pos < chunk_end:
                    tables_in_chunk.append((table_pos - chunk_start, table_content))

            # If no tables in this chunk, add it as is
            if not tables_in_chunk:
                final_chunks.append(chunk)
                continue

            # Insert tables at correct positions
            modified_chunk = chunk
            offset = 0
            for position, table_content in sorted(tables_in_chunk):
                position += offset  # Adjust for previously inserted tables
                modified_chunk = modified_chunk[:position] + "\n\n" + table_content + "\n\n" + modified_chunk[position:]
                offset += len("\n\n" + table_content + "\n\n")

            final_chunks.append(modified_chunk)

        for i, chunk in enumerate(final_chunks):
            print(f"\n\n--- Chunk {i + 1} ---")
            print(chunk)

        print(f"\nChunked text numbers: {len(final_chunks)}")
        return final_chunks





if __name__ == "__main__":
    extractor = DocxTextExtractor(chunk_size=1024, chunk_overlap=128)
    docx_path = "data/3-page.docx"

    documents = extractor.extract_content(docx_path)

    output_file = "data/chunked_texts_docx.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(documents, f, ensure_ascii=False, indent=4)

